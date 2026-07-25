import json
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def process_folder_simple(input_folder, output_folder='output'):
    """
    Process all JSON files without using textract-trp
    """
    # Create output folder
    os.makedirs(output_folder, exist_ok=True)
    
    # Get all JSON files
    json_files = [f for f in os.listdir(input_folder) if f.endswith('.json')]
    
    print(f"Found {len(json_files)} JSON files")
    
    for idx, json_file in enumerate(json_files, 1):
        try:
            json_path = os.path.join(input_folder, json_file)
            
            print(f"[{idx}/{len(json_files)}] Processing: {json_file}")
            
            # Read JSON
            with open(json_path, 'r', encoding='utf-8') as f:
                response = json.load(f)
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_paragraph(f"Document: {json_file}")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.bold = True
            title.runs[0].font.size = Pt(14)
            doc.add_paragraph()  # spacing
            
            # Track pages
            current_page = None
            line_count = 0
            
            # Process blocks
            for block in response.get('Blocks', []):
                if block['BlockType'] == 'LINE':
                    # Add page break when page changes
                    page_num = block.get('Page', 1)
                    if current_page is not None and page_num != current_page:
                        doc.add_page_break()
                        # Add page header
                        page_header = doc.add_paragraph(f"--- Page {page_num} ---")
                        page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    current_page = page_num
                    
                    # Add line text
                    para = doc.add_paragraph(block['Text'])
                    for run in para.runs:
                        run.font.size = Pt(11)
                    
                    line_count += 1
            
            # Save
            output_name = os.path.splitext(json_file)[0] + '.docx'
            output_path = os.path.join(output_folder, output_name)
            doc.save(output_path)
            
            print(f"  ✓ Saved: {output_name}")
            print(f"    Lines extracted: {line_count}")
            
        except Exception as e:
            print(f"  ✗ Error: {str(e)}")

# Usage
if __name__ == "__main__":
    input_folder = r'C:\Users\orl12\OneDrive\Documentos\6. Quality Management\My Procedures\textract_parser\Product Quality Complaint Management'      # ← Change to your folder path
    output_folder = r'C:\Users\orl12\OneDrive\Documentos\6. Quality Management\My Procedures\textract_parser\Product Quality Complaint Management'       # ← Change to your output folder
    
    process_folder_simple(input_folder, output_folder)